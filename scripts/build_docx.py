"""Generate manuscript DOCX using narrative_proposal preset."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor

MANUSCRIPT = Path(r"F:\ONC\manuscript\manuscript.md")
OUTPUT = Path(r"F:\ONC\manuscript\manuscript.docx")
BODY_FONT = "Calibri"
BODY_SIZE = 11

def sanitize(t):
    return "".join(ch for ch in t if ord(ch) >= 32 or ord(ch) in (9, 10, 13))

def setup_styles(doc):
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(BODY_SIZE)
    style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15
    for lvl, sz, clr in [(1, 16, RGBColor(0x2E,0x74,0xB5)), (2, 13, RGBColor(0x2E,0x74,0xB5)), (3, 12, RGBColor(0x1F,0x4D,0x78))]:
        s = doc.styles[f"Heading {lvl}"]
        s.font.name = BODY_FONT
        s.font.size = Pt(sz)
        s.font.color.rgb = clr
        s.font.bold = True
        s.paragraph_format.space_before = Pt({1:24,2:18,3:12}[lvl])
        s.paragraph_format.space_after = Pt({1:8,2:6,3:4}[lvl])
        s.paragraph_format.line_spacing = 1.15
    for sec in doc.sections:
        sec.top_margin = Inches(1.0)
        sec.right_margin = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin = Inches(1.0)

def build_doc(md_path):
    doc = Document()
    setup_styles(doc)
    text = sanitize(Path(md_path).read_text(encoding="utf-8"))
    lines = text.split("\n")
    in_code = False
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("`"):
            in_code = not in_code
            continue
        if in_code:
            p = doc.add_paragraph(stripped)
            p.style = doc.styles["Normal"]
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.0
            for run in p.runs:
                run.font.name = "Consolas"
                run.font.size = Pt(9)
            continue
        if not stripped:
            continue
        if stripped == "---":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pPr = p._p.get_or_add_pPr()
            from docx.oxml.ns import qn
            pBdr = pPr.makeelement(qn("w:pBdr"), {})
            bottom = pBdr.makeelement(qn("w:bottom"), {qn("w:val"):"single",qn("w:sz"):"4",qn("w:space"):"4",qn("w:color"):"999999"})
            pBdr.append(bottom)
            pPr.append(pBdr)
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=1)
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2)
            continue
        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=3)
            continue
        if stripped.startswith("**") and stripped.count("**") >= 2:
            end = stripped.index("**", 2)
            bold_text = stripped[2:end]
            rest = stripped[end+2:].strip()
            p = doc.add_paragraph()
            r = p.add_run(sanitize(bold_text))
            r.bold = True
            r.font.name = BODY_FONT
            r.font.size = Pt(BODY_SIZE)
            if rest:
                r2 = p.add_run(" " + sanitize(rest))
                r2.font.name = BODY_FONT
                r2.font.size = Pt(BODY_SIZE)
            p.paragraph_format.space_before = Pt(8)
            continue
        p = doc.add_paragraph()
        parts = re.split(r"(\*\*.+?\*\*)", stripped)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                r = p.add_run(sanitize(part[2:-2]))
                r.bold = True
            else:
                r = p.add_run(sanitize(part))
            r.font.name = BODY_FONT
            r.font.size = Pt(BODY_SIZE)
    return doc

if __name__ == "__main__":
    doc = build_doc(MANUSCRIPT)
    doc.save(str(OUTPUT))
    print(f"DOCX saved: {OUTPUT}")
